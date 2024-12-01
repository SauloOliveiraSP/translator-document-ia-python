import requests
from docx import Document
import os

# Configurações do Serviço Azure Translator
subscription_key = "sua_subscription_key"
endpoint = 'https://api.cognitive.microsofttranslator.com'
location = "eastus2"
language_destination = 'pt-br'  # Idioma desejado para tradução

# Função para traduzir texto
def translator_text(text, target_language):
    path = '/translate'
    constructed_url = endpoint + path
    headers = {
        'Ocp-Apim-Subscription-Key': subscription_key,
        'Ocp-Apim-Subscription-Region': location,
        'Content-type': 'application/json',
        'x-ClientTraceId': str(os.urandom(16))
    }
    body = [{'text': text}]
    params = {'api-version': "3.0", 'from': 'en', 'to': target_language}
    request = requests.post(constructed_url, params=params, headers=headers, json=body)
    response = request.json()
    return response[0]["translations"][0]["text"]

# Função para traduzir documentos .docx
def traslate_document(path):
    document = Document(path)
    full_text = []

    for paragraph in document.paragraphs:
        translated_text = translator_text(paragraph.text, language_destination)
        full_text.append(translated_text)

    translated_doc = Document()
    for line in full_text:
        translated_doc.add_paragraph(line)
    
    path_translated = path.replace(".docx", f"_{language_destination}.docx")
    translated_doc.save(path_translated)
    return path_translated

# Exemplo de Uso
translator_text("Hello World!", language_destination)  # Traduz texto simples
input_file = "musica.docx"  # Substitua pelo caminho do arquivo .docx
traslate_document(input_file)  # Traduz o documento
